from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Pt


TEMPLATE = Path(r"C:\Users\ahmad\Downloads\Birketts_Junior_Data_Engineer_CV.docx")
OUT = Path("output/pdf/Ahmad_Bokhari_CV_Fitch_Data_Analytics_Senior_Associate.docx")


def clear_body(document: Document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(paragraph, size: float | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Aptos"
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def para(doc: Document, text: str = "", style: str = "Normal", size: float = 9.3, bold: bool = False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(2.2)
    p.paragraph_format.line_spacing = 1.02
    return p


def heading(doc: Document, text: str):
    p = para(doc, text.upper(), size=9.6, bold=True)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    return p


def bullet(doc: Document, text: str):
    p = para(doc, text, style="List Paragraph", size=8.75)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.first_line_indent = Pt(-7)
    return p


def role_line(doc: Document, title: str, date: str = ""):
    p = para(doc, size=8.9, bold=True)
    p.add_run(title).bold = True
    if date:
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Pt(468), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t" + date)
    set_run_font(p, 8.9)
    return p


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Pt(35)
    section.bottom_margin = Pt(35)
    section.left_margin = Pt(44)
    section.right_margin = Pt(44)

    name = para(doc, "Syed Muhammad Ahmad Bokhari", size=16, bold=True)
    name.paragraph_format.space_after = Pt(0)
    para(doc, "Data Analytics / AI Engineer | LLM Workflows, RAG, Python and Analytics Prototyping", size=10.2)
    para(
        doc,
        "Bradford, UK | 07769685019 | bokhariahmed765@gmail.com | "
        "linkedin.com/in/syed-muhammad-ahmad-bokhari-aa5064294 | github.com/syedahmadbokhari",
        size=8.35,
    )

    heading(doc, "Personal Summary")
    para(
        doc,
        "First-Class BSc Applied Artificial Intelligence graduate with hands-on project evidence in "
        "LLM-powered analytics, RAG, agentic AI workflows, evaluation, data pipelines and BI reporting. "
        "Built a deployed FastAPI/React RAG assistant over real portfolio documentation using "
        "MiniLM embeddings, FAISS retrieval, Groq generation and an Anthropic tool-use agent path, with "
        "deterministic grounding assertions and trace logging. Strong Python and SQL fundamentals, "
        "comfortable turning ambiguous ideas into tested prototypes and explaining trade-offs to "
        "technical and non-technical stakeholders.",
        size=8.85,
    )

    heading(doc, "Education")
    role_line(doc, "BSc (Hons) Applied Artificial Intelligence - University of Bradford", "2023 - 2026")
    para(doc, "First Class Honours", size=8.7)

    heading(doc, "Role-Relevant Skills")
    bullet(doc, "LLM and RAG systems: markdown-section chunking, sentence-transformer embeddings, FAISS vector search, source-grounded generation, refusal gates, ReAct-style tool-use agents and deterministic post-generation validation.")
    bullet(doc, "Python engineering: readable, tested Python; FastAPI; Pydantic; JWT authentication; rate limiting; pytest; package extraction; CLI workflows; JSON/CSV evaluation artifacts.")
    bullet(doc, "Analytics and data engineering: SQL, dbt, Airflow, DuckDB, PostgreSQL, BigQuery dry-run analysis, Snowflake integration code, AWS S3, Kafka demo pipeline, Docker and Docker Compose.")
    bullet(doc, "BI and communication: Power BI, Tableau, Streamlit, Excel analysis workbooks, KPI reporting, plain-language summaries and stakeholder-facing documentation.")
    bullet(doc, "Quality and delivery: Git, GitHub Actions CI, data quality gates, Great Expectations, dbt tests, regression tests, live validation reports and pragmatic trade-off documentation.")

    heading(doc, "Selected Projects")
    role_line(doc, "AI Interview Prep Assistant - LLM Analytics and Agentic RAG System", "GitHub: syedahmadbokhari/ai-interview-prep")
    para(doc, "Python, FastAPI, React, FAISS, sentence-transformers, Groq, Anthropic tool_use, Docker, GitHub Actions", size=8.35)
    bullet(doc, "Built a deployed project-knowledge assistant that answers natural-language questions over real GitHub project documentation with grounded citations and an honest no-result path before any LLM call.")
    bullet(doc, "Implemented a RAG pipeline with 70 markdown-section chunks, local MiniLM embeddings, exact FAISS cosine search and a measured relevance threshold; the original 12-case evaluation recorded 10/10 grounded retrieval and 2/2 off-topic refusals.")
    bullet(doc, "Added a semi-autonomous Anthropic tool-use agent with list/search/project-summary tools, JSONL traces, maximum-iteration handling and deterministic assertion validation for entities, technologies, metrics, dates and scope.")
    bullet(doc, "Wrapped the system in a JWT-protected FastAPI API and React frontend, with Docker deployment on Render and CI covering tests, package build, deterministic mock evaluation and an ML quality gate.")

    role_line(doc, "Retail Data Platform - Analytics Marts and BI Reporting")
    para(doc, "SQL, Python, dbt, Airflow, PostgreSQL, BigQuery, Snowflake code, Power BI, Tableau, Streamlit, Excel", size=8.35)
    bullet(doc, "Built a production-style retail analytics platform that turns raw sales events into layered clean and analytics tables for revenue, product, discount and traffic reporting.")
    bullet(doc, "Created BI outputs in Power BI, Tableau, Streamlit and a native Excel workbook, including stakeholder-readable summaries and auditable formulas for pricing opportunity modelling.")
    bullet(doc, "Detected and excluded 354 corrupted GBP 0 price records before feature engineering, preventing downstream recommendation and reporting outputs from being distorted.")
    bullet(doc, "Used statistical testing to validate discount/revenue patterns and fixed a misleading BigQuery date-window bug before reporting the verified 58.7% bytes-scanned reduction from partitioning/clustering.")

    role_line(doc, "UK Crime Data Pipeline - Cloud Ingestion, Data Quality and Geospatial Analytics")
    para(doc, "Python, AWS S3, DuckDB, dbt, Airflow, FastAPI, Great Expectations, Kafka, Streamlit, Terraform, GitHub Actions", size=8.35)
    bullet(doc, "Designed an end-to-end UK Police open-data pipeline covering ingestion, Hive-style S3 partitioning, DuckDB warehousing, dbt staging/marts, orchestration and dashboard delivery.")
    bullet(doc, "Modelled analytics marts for crime categories, monthly trends, force outcomes and LSOA hotspots, supporting operational questions such as trend, geography and outcome-rate analysis.")
    bullet(doc, "Formalised data quality with Great Expectations checks for required columns, nulls, accepted values, UK coordinate bounds and non-null crime ID uniqueness, alongside dbt tests and DAG row-count gates.")
    bullet(doc, "Exposed outputs through a Streamlit/Folium dashboard and FastAPI service, with Docker Compose setup and GitHub Actions running pytest plus dbt compile.")

    heading(doc, "Work Experience")
    role_line(doc, "AI Project Lead, Bradford Council", "Jul 2026 - Present")
    bullet(doc, "Leading delivery of a data-driven AI detection and monitoring system for a local council, translating model outputs into reporting that operational stakeholders can act on.")
    bullet(doc, "Coordinating requirements and reporting cadence with Council stakeholders within governance and data-protection constraints.")
    bullet(doc, "Leading a team of three across an approximately six-month engagement, using fortnightly reporting to manage priorities, risks and progress.")

    role_line(doc, "Data Engineer Intern, Holy Family Hospital, Pakistan", "Apr 2023 - Jun 2023")
    bullet(doc, "Extracted, cleaned and organised patient and operational datasets from hospital databases to support reporting and analysis.")
    bullet(doc, "Wrote SQL queries to retrieve, filter and validate records, helping identify missing or inconsistent data entries.")
    bullet(doc, "Supported Python ETL scripts for repetitive data entry and formatting tasks, while maintaining confidentiality in healthcare data handling.")

    heading(doc, "Technical Toolkit")
    para(doc, "Python, SQL, FastAPI, Pydantic, Pandas, NumPy, pytest, sentence-transformers, FAISS, Anthropic SDK, Groq, REST APIs, JWT, Docker, Docker Compose, Git, GitHub Actions", size=8.55)
    para(doc, "dbt, Airflow, PostgreSQL, DuckDB, BigQuery, Snowflake integration code, AWS S3, Kafka, Great Expectations, Power BI, Tableau, Streamlit, Excel", size=8.55)

    heading(doc, "Certifications")
    para(doc, "IBM SkillsBuild - Data Fundamentals and AI Fundamentals Certificates (2026)", size=8.7)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
