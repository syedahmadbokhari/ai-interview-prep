from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output/pdf/Ahmad_CV_MHA_DataAnalyst_Template.docx")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_pr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_pr.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def set_para_border_bottom(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def set_run(run, bold=False, italic=False, size=10, color="222222"):
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_mixed_para(doc, parts, style_name=None, space_after=4, left_indent=None):
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    for text, bold, italic in parts:
        r = p.add_run(text)
        set_run(r, bold=bold, italic=italic)
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title.upper())
    set_run(r, bold=True, size=11.5, color="000000")
    set_para_border_bottom(p)


def add_bullet(doc, parts):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.15)
    p.paragraph_format.space_after = Pt(3.5)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, italic in parts:
        r = p.add_run(text)
        set_run(r, bold=bold, italic=italic, size=9.6)
    return p


def add_project(doc, title, tech, bullets):
    add_mixed_para(doc, [(title, True, False)], space_after=1)
    add_mixed_para(doc, [(tech, False, True)], space_after=2)
    for item in bullets:
        add_bullet(doc, item)


def add_two_col(doc, left, right):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(13.2)
    table.columns[1].width = Cm(4.3)
    for cell in table.row_cells(0):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(
            cell,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
    left_p = table.cell(0, 0).paragraphs[0]
    right_p = table.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lr = left_p.add_run(left)
    rr = right_p.add_run(right)
    set_run(lr, bold=True)
    set_run(rr, italic=True)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(8)
    r = name.add_run("SYED MUHAMMAD AHMAD BOKHARI")
    set_run(r, bold=True, size=18, color="2F3744")

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_after = Pt(8)
    r = role.add_run("Data Analyst | SQL, Power BI & Insight-Led Storytelling")
    set_run(r, size=12, color="333333")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(12)
    r = contact.add_run(
        "07769685019 | bokhariahmed765@gmail.com | linkedin.com/in/syed-muhammad-ahmad-bokhari-aa5064294\n"
        "github.com/syedahmadbokhari"
    )
    set_run(r, size=8.8, color="3F3F3F")

    add_section(doc, "Personal Summary")
    add_mixed_para(
        doc,
        [
            (
                "AI & Data Professional with a First-Class degree from the University of Bradford. I use SQL, Power BI, Excel, Python and dbt to turn raw data into clear insight, dashboards and practical recommendations. Project experience includes data transformation, KPI reporting, service-focused analysis, validation checks and insight-led storytelling for technical and non-technical stakeholders, with a strong interest in using data to improve services and community outcomes.",
                False,
                False,
            )
        ],
    )

    add_section(doc, "Education")
    add_two_col(doc, "BSc (Hons) Applied Artificial Intelligence - University of Bradford", "2023 - 2026")
    add_mixed_para(doc, [("First Class Honours", False, False)], space_after=4)

    add_section(doc, "Projects")
    add_project(
        doc,
        "UK Crime Data Pipeline - Public Service Insight & Reporting",
        "Python, SQL, DuckDB, dbt, Airflow, FastAPI, Great Expectations, Streamlit",
        [
            [("Built public-service dashboards", True, False), (" for crime trends, category breakdowns, geographic hotspots and force KPIs, turning complex operational data into insight stakeholders could act on.", False, False)],
            [("Created an end-to-end data workflow", True, False), (" from source CSV ingestion through DuckDB warehouse tables, dbt staging/marts, Airflow orchestration and Streamlit reporting.", False, False)],
            [("Automated data quality checks", True, False), (" for schema, nulls, accepted values, row counts, coordinate ranges and unique IDs so reporting outputs stayed accurate and trustworthy.", False, False)],
            [("Used insight-led storytelling", True, False), (" to explain trends, risks, outcome rates and hotspot patterns in plain language for non-technical users.", False, False)],
        ],
    )
    add_project(
        doc,
        "Retail Data Platform - Power BI, Excel & Commercial Insight",
        "SQL, Python, PostgreSQL, dbt, Airflow, Power BI, Tableau, Excel, Snowflake, Docker",
        [
            [("Designed Power BI, Tableau and Excel reports", True, False), (" covering revenue, discount impact, traffic trends, product performance and pricing opportunities for decision-making.", False, False)],
            [("Modelled clean reporting datasets", True, False), (" across raw, clean and analytics layers using SQL and dbt, creating repeatable KPI tables for dashboard use.", False, False)],
            [("Used SQL analysis", True, False), (" with CTEs, window functions, joins and aggregations to identify performance trends, risks and opportunities.", False, False)],
            [("Improved reporting confidence", True, False), (" by identifying 354 corrupted GBP 0 price records before feature engineering and downstream dashboard use.", False, False)],
        ],
    )
    add_project(
        doc,
        "AI for Environmental Monitoring & Urban Planning - BSc final-year project",
        "Python, OpenCV, PyTorch (MobileNetV3), MOG2, ByteTrack, YOLOv8, Flask, GIS",
        [
            [("Translated AI outputs into a service workflow", True, False), (" with alerts, snapshots, confidence scores and GIS hotspot mapping for practical stakeholder review.", False, False)],
            [("Evaluated the system against manually observed ground truth", True, False), (", reporting 75% accuracy and 60% recall while explaining model tradeoffs clearly.", False, False)],
            [("Designed a feedback-led improvement process", True, False), (" where dismissed alerts become retraining data, supporting continuous improvement and better future performance.", False, False)],
        ],
    )

    doc.add_page_break()
    add_section(doc, "Work Experience")
    add_two_col(doc, "AI Project Lead, Bradford Council", "Jul 2026 - Present")
    add_bullet(doc, [("Leading a data-driven AI monitoring project", True, False), (" for a local council, translating a detection system into reporting stakeholders can act on", False, False)])
    add_bullet(doc, [("Coordinating requirements and reporting cadence with Council stakeholders within governance and data protection constraints", False, False)])
    add_bullet(doc, [("Leading a team of three across a ~6-month engagement, with fortnightly reporting to Council and academic stakeholders", False, False)])

    add_section(doc, "Technical Skills")
    skills = [
        [("SQL & Data Analysis:", True, False), (" SQL, CTEs, window functions, multi-table joins, trend analysis, performance insight, risk/opportunity identification", False, False)],
        [("Power BI & Reporting:", True, False), (" Power BI, DAX concepts, Excel, Tableau, Streamlit, KPI dashboards, self-serve reporting, insight-led storytelling", False, False)],
        [("Data Transformation:", True, False), (" Python, Pandas, NumPy, dbt, PostgreSQL, DuckDB, ETL pipelines, clean analytics layers, data workflow mapping", False, False)],
        [("Data Quality:", True, False), (" validation gates, schema testing, accepted-value checks, null handling, reconciliation, automated checks, auditable assumptions", False, False)],
        [("Digital & Service Improvement:", True, False), (" dashboard adoption, process improvement, customer-focused insight, community/public-sector data, continuous improvement", False, False)],
        [("Collaboration:", True, False), (" requirements gathering, stakeholder reporting, technical/non-technical communication, Git/GitHub, documentation, team working", False, False)],
    ]
    for line in skills:
        add_mixed_para(doc, line, space_after=4)

    add_section(doc, "Certifications")
    add_bullet(doc, [("Data Fundamentals & AI Fundamentals Certificates, IBM SkillsBuild (2026)", False, False)])

    add_section(doc, "Soft Skills")
    add_mixed_para(doc, [("Problem-Solving, Teamwork, Communication, Stakeholder Reporting.", False, False)])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
